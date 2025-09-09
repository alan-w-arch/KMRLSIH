# app/services/text_extractor.py
import io
import logging
from typing import Optional, Dict, Any
import asyncio

# PDF extraction
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not installed - PDF text extraction disabled")

# DOCX extraction  
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed - DOCX text extraction disabled")

# XLSX extraction
try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    logging.warning("openpyxl not installed - XLSX text extraction disabled")

logger = logging.getLogger(__name__)


async def extract_text_from_file(
    data: bytes, 
    content_type: str,
    filename: str = None
) -> Optional[Dict[str, Any]]:
    """
    Extract text from various file types.
    Returns dict with extracted text and metadata, or None if not supported.
    """
    try:
        if content_type == "application/pdf":
            return await _extract_pdf_text(data)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await _extract_docx_text(data)
        elif content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return await _extract_xlsx_text(data)
        elif content_type == "text/plain":
            return await _extract_plain_text(data)
        else:
            logger.debug(f"Text extraction not supported for content type: {content_type}")
            return None
            
    except Exception as e:
        logger.exception(f"Text extraction failed for {filename}: {e}")
        return None


async def _extract_pdf_text(data: bytes) -> Optional[Dict[str, Any]]:
    """Extract text from PDF using PyPDF2."""
    if not PDF_AVAILABLE:
        return None
    
    def _extract():
        try:
            pdf_file = io.BytesIO(data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            # Get metadata if available
            metadata = {}
            if pdf_reader.metadata:
                metadata = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                }
            
            return {
                'text': text_content.strip(),
                'page_count': page_count,
                'extraction_method': 'pdf_direct',
                'metadata': metadata
            }
            
        except Exception as e:
            logger.exception(f"PDF text extraction failed: {e}")
            return None
    
    return await asyncio.to_thread(_extract)


async def _extract_docx_text(data: bytes) -> Optional[Dict[str, Any]]:
    """Extract text from DOCX using python-docx."""
    if not DOCX_AVAILABLE:
        return None
    
    def _extract():
        try:
            docx_file = io.BytesIO(data)
            doc = DocxDocument(docx_file)
            
            # Extract paragraph text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract table text
            tables = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    tables.append("\n".join(table_text))
            
            # Combine all text
            all_text = []
            if paragraphs:
                all_text.append("--- Document Text ---\n" + "\n\n".join(paragraphs))
            if tables:
                all_text.append("--- Tables ---\n" + "\n\n".join(tables))
            
            # Get document properties
            properties = {}
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                properties['title'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                properties['author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                properties['subject'] = doc.core_properties.subject
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                properties['created'] = str(doc.core_properties.created)
            
            return {
                'text': "\n\n".join(all_text),
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'extraction_method': 'docx_direct',
                'metadata': properties
            }
            
        except Exception as e:
            logger.exception(f"DOCX text extraction failed: {e}")
            return None
    
    return await asyncio.to_thread(_extract)


async def _extract_xlsx_text(data: bytes) -> Optional[Dict[str, Any]]:
    """Extract text from XLSX using openpyxl."""
    if not XLSX_AVAILABLE:
        return None
    
    def _extract():
        try:
            xlsx_file = io.BytesIO(data)
            workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
            
            all_sheets_text = []
            sheet_info = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                # Get all non-empty cells
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                    if row_text:
                        sheet_text.append(" | ".join(row_text))
                
                if sheet_text:
                    sheet_content = f"--- Sheet: {sheet_name} ---\n" + "\n".join(sheet_text)
                    all_sheets_text.append(sheet_content)
                    sheet_info.append({
                        'name': sheet_name,
                        'row_count': len(sheet_text)
                    })
            
            # Get workbook properties
            properties = {}
            if hasattr(workbook.properties, 'title') and workbook.properties.title:
                properties['title'] = workbook.properties.title
            if hasattr(workbook.properties, 'creator') and workbook.properties.creator:
                properties['creator'] = workbook.properties.creator
            if hasattr(workbook.properties, 'created') and workbook.properties.created:
                properties['created'] = str(workbook.properties.created)
            
            return {
                'text': "\n\n".join(all_sheets_text),
                'sheet_count': len(all_sheets_text),
                'sheets': sheet_info,
                'extraction_method': 'xlsx_direct',
                'metadata': properties
            }
            
        except Exception as e:
            logger.exception(f"XLSX text extraction failed: {e}")
            return None
    
    return await asyncio.to_thread(_extract)


async def _extract_plain_text(data: bytes) -> Optional[Dict[str, Any]]:
    """Extract text from plain text files."""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                text = data.decode(encoding)
                return {
                    'text': text,
                    'encoding': encoding,
                    'extraction_method': 'plain_text',
                    'metadata': {'character_count': len(text)}
                }
            except UnicodeDecodeError:
                continue
        
        logger.warning("Failed to decode text file with any supported encoding")
        return None
        
    except Exception as e:
        logger.exception(f"Plain text extraction failed: {e}")
        return None